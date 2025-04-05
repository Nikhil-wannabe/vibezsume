def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file"""
    try:
        import docx
        
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except ImportError:
        return "Error: DOCX parsing library (python-docx) not installed."
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"