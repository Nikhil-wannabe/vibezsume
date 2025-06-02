"""
This module provides utility functions for extracting plain text content
from various file formats, commonly used for resumes.
It currently supports PDF, DOCX, and plain TXT files.
"""
import PyPDF2
from PyPDF2.errors import FileNotDecryptedError # For specific encrypted PDF error handling
import docx
from typing import Optional # Union was not used, so removed.
import io

def extract_text_from_pdf(file_content: bytes) -> Optional[str]:
    """
    Extracts all text content from a PDF file's byte content.

    Args:
        file_content (bytes): The byte content of the PDF file.

    Returns:
        Optional[str]: The concatenated text extracted from all pages of the PDF.
                       Returns None if extraction fails (e.g., corrupted file),
                       the PDF is encrypted, or if no meaningful text could be extracted.

    Limitations:
    - **Extraction Quality**: Quality of text extraction can vary significantly for complex PDFs,
      especially those with multi-column layouts, tables, unconventional fonts, embedded graphics
      with text overlays, or rotated text. It's best suited for straightforward text documents.
    - **OCR Not Included**: This function does not perform Optical Character Recognition (OCR).
      Text embedded within images or from scanned (image-based) PDFs will not be extracted.
      An external OCR tool would be necessary for such cases.
    - **Encryption**: Encrypted or password-protected PDFs will typically cause extraction to fail,
      though a specific check for `FileNotDecryptedError` is included to provide a clearer message.
    - **PyPDF2 Limitations**: Relies on the PyPDF2 library. The success and fidelity of extraction
      are subject to PyPDF2's capabilities and any inherent bugs or limitations it may have with
      certain PDF structures or versions.
    """
    try:
        # Create a file-like object from the byte content to be read by PyPDF2.
        pdf_file = io.BytesIO(file_content)
        # Initialize a PDF reader object.
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = [] # List to hold text extracted from each page.
        # Iterate through each page of the PDF.
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                # Extract text from the current page.
                page_text = page.extract_text()
                # Ensure page_text is not None and contains more than just whitespace before adding.
                if page_text and page_text.strip(): 
                    text_parts.append(page_text)
            except Exception as e_page:
                # Log error for a specific page but attempt to continue with other pages.
                # This makes the extraction more resilient to single-page failures.
                print(f"Error extracting text from PDF page {page_num + 1}: {e_page}")
                continue # Try next page

        # Join text from all pages and strip leading/trailing whitespace from the final result.
        full_text = "".join(text_parts).strip()
        # Return the extracted text only if it's not empty, otherwise None.
        return full_text if full_text else None 
    except FileNotDecryptedError:
        # Specific handling for encrypted PDFs that PyPDF2 cannot open.
        print("Error: Could not extract text from PDF. The file is encrypted and password-protected.")
        return None
    except Exception as e:
        # General error handling for other PyPDF2 issues (e.g., malformed PDF, I/O errors) or unexpected errors.
        print(f"Error extracting text from PDF: {e}")
        return None

def extract_text_from_docx(file_content: bytes) -> Optional[str]:
    """
    Extracts all text content from a DOCX file's byte content.

    Args:
        file_content (bytes): The byte content of the DOCX file.

    Returns:
        Optional[str]: The concatenated text extracted from all paragraphs in the DOCX file.
                       Returns None if extraction fails or no text is found.

    Limitations:
    - May not perfectly preserve complex formatting (e.g., tables, headers/footers text boxes).
      It primarily extracts text from main document paragraphs.
    - Images or embedded objects containing text will not have their text extracted.
    """
    try:
        # Create a file-like object from the byte content for python-docx.
        doc_file = io.BytesIO(file_content)
        # Load the DOCX document.
        doc_obj = docx.Document(doc_file)
        
        # Extract text from each paragraph and join them with newlines.
        text = "\n".join(paragraph.text for paragraph in doc_obj.paragraphs)
        # Strip leading/trailing whitespace from the final text.
        return text.strip() if text.strip() else None
    except Exception as e:
        # Handle errors during DOCX parsing or text extraction.
        print(f"Error extracting text from DOCX: {e}")
        return None

def extract_text_from_file(file_name: str, file_content: bytes) -> Optional[str]:
    """
    Extracts text from a file based on its extension, supporting PDF, DOCX, and TXT.

    Args:
        file_name (str): The name of the file (e.g., "resume.pdf"). This is used
                         to determine the file type by its extension.
        file_content (bytes): The byte content of the file.

    Returns:
        Optional[str]: The extracted text as a string if successful and the format is supported,
                       otherwise None.
    """
    if not file_name or not file_content:
        # Basic validation for inputs.
        print("Error: File name or content is missing for text extraction.")
        return None

    # Determine file extension to select the appropriate extraction method.
    file_extension = file_name.split('.')[-1].lower()

    if file_extension == 'pdf':
        return extract_text_from_pdf(file_content)
    elif file_extension == 'docx':
        return extract_text_from_docx(file_content)
    elif file_extension == 'txt':
        try:
            # Decode as UTF-8 first, as it's most common.
            return file_content.decode('utf-8').strip()
        except UnicodeDecodeError:
            # If UTF-8 fails, try latin-1 as a common fallback.
            try:
                return file_content.decode('latin-1').strip()
            except Exception as e_decode:
                print(f"Error decoding TXT file with UTF-8 and latin-1: {e_decode}")
                return None
    else:
        # Handle unsupported file formats.
        print(f"Unsupported file format for text extraction: '{file_extension}'")
        return None

if __name__ == '__main__':
    # Basic test logic (requires dummy files or mocking for thorough testing).
    # Create dummy PDF and DOCX for local testing if possible, or use mocks.
    # For now, just print a message.
    print("Text extraction module ready. Testing requires actual file content.")

    # Example of how it might be used (conceptual)
    # with open("dummy.pdf", "rb") as f:
    #     content = f.read()
    #     text = extract_text_from_file("dummy.pdf", content)
    #     if text:
    #         print("Successfully extracted text from PDF (dummy):", text[:100])
    #     else:
    #         print("Failed to extract text from PDF (dummy).")

    # try:
    #     # Create a dummy docx for testing
    #     from docx import Document
    #     doc = Document()
    #     doc.add_paragraph("This is a test document for DOCX extraction.")
    #     doc.add_paragraph("It contains multiple paragraphs.")
    #     dummy_docx_bytes_io = io.BytesIO()
    #     doc.save(dummy_docx_bytes_io)
    #     dummy_docx_bytes = dummy_docx_bytes_io.getvalue()

    #     text_docx = extract_text_from_file("test.docx", dummy_docx_bytes)
    #     if text_docx:
    #         print(f"Successfully extracted from dummy DOCX: '{text_docx}'")
    #     else:
    #         print("Failed to extract from dummy DOCX.")

    #     # Create dummy txt for testing
    #     dummy_txt_content = "This is a plain text file for testing. With UTF-8 characters: éàçüö".encode('utf-8')
    #     text_txt = extract_text_from_file("test.txt", dummy_txt_content)
    #     if text_txt:
    #         print(f"Successfully extracted from dummy TXT: '{text_txt}'")
    #     else:
    #         print("Failed to extract from dummy TXT.")

    # except ImportError:
    #     print("python-docx or PyPDF2 not installed, skipping some local tests.")
    # except Exception as e:
    #     print(f"Error during local test setup: {e}")
